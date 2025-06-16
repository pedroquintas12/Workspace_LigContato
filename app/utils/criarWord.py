import io
import requests
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def gerar_relatorio_word(relatorios):
    doc = Document()

    # Adiciona a imagem no topo centralizada
    imagem_url = "https://resources-lig.s3.amazonaws.com/images/logo_thumbnail.png"
    response = requests.get(imagem_url)
    if response.status_code == 200:
        imagem_stream = io.BytesIO(response.content)
        doc.add_picture(imagem_stream, width=Inches(1.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for idx, relatorio in enumerate(relatorios):
        codigo = relatorio["codigo"]
        dados = relatorio["dados"]

        # Título para cada código
        doc.add_paragraph()
        p_codigo = doc.add_paragraph(f"Relatório do Escritório Código: {codigo}")
        p_codigo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run_codigo = p_codigo.runs[0]
        run_codigo.bold = True
        run_codigo.font.size = Pt(16)
        run_codigo.underline = True

        doc.add_paragraph()  # Espaço antes dos clientes

        if not dados:
            p = doc.add_paragraph("Nenhum dado disponível.")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            nome_escritorio = dados[0]['nome_escritorio']
            p_escritorio = doc.add_paragraph()
            p_escritorio.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run_esc = p_escritorio.add_run(f"Escritório: {nome_escritorio}")
            run_esc.bold = True
            run_esc.font.size = Pt(14)
            run_esc.underline = True

            doc.add_paragraph()  # Espaço antes dos clientes

            for item in dados:
                p_cliente = doc.add_paragraph()
                p_cliente.paragraph_format.space_after = Pt(2)
                run_cliente = p_cliente.add_run(f"{item['nome_cliente']}")
                run_cliente.bold = True
                run_cliente.font.size = Pt(12)

                if item['abrangencia_por_estado']:
                    p_abrangencia_title = doc.add_paragraph()
                    run_abrangencia = p_abrangencia_title.add_run("Abrangência por estado:")
                    run_abrangencia.bold = True
                    run_abrangencia.font.size = Pt(11)

                    for estado, abrangencias_estado in item['abrangencia_por_estado'].items():
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Pt(24)
                        run_estado = p.add_run(f"{estado}: ")
                        run_estado.bold = True
                        run_estado.font.size = Pt(11)

                        run_abr = p.add_run(', '.join(abrangencias_estado))
                        run_abr.font.size = Pt(10)
                else:
                    p_nao_abrangencia = doc.add_paragraph("Abrangência por estado: Nenhuma")
                    p_nao_abrangencia.paragraph_format.left_indent = Pt(12)
                    run = p_nao_abrangencia.runs[0]
                    run.font.size = Pt(10)
                    run.italic = True

                # Linha horizontal separadora
                hr = doc.add_paragraph()
                hr.add_run("─" * 40).italic = True
                hr.paragraph_format.space_before = Pt(6)
                hr.paragraph_format.space_after = Pt(6)
                hr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Adiciona uma quebra de página entre relatórios, exceto no último
        if idx < len(relatorios) - 1:
            doc.add_page_break()

    arquivo_memoria = io.BytesIO()
    doc.save(arquivo_memoria)
    arquivo_memoria.seek(0)
    return arquivo_memoria
